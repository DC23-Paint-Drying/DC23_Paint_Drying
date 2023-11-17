import pytest
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import qn

import src.report.report_utils as utils


data = {
    'timestamp': '2023-11-13',
    'date': '2023-11',
    'users': {
        'subscriptions': {'Podstawowy': 0, 'Standardowy': 0, 'Premium': 0},
        'gender': {'Mężczyźni': 0, 'Kobiety': 0, 'Inne': 0},
        'age': {'<18': 0, '18-25': 0, '26-35': 0, '36-50': 0, '>50': 0}
    },
    'sales': {
        'Podstawowy': {'users': 0, 'profit': 0.0},
        'Standardowy': {'users': 0, 'profit': 0.0},
        'Premium': {'users': 0, 'profit': 0.0}
    },
    'recent': {
        'users': 0,
        'subscribed': {'Podstawowy': 0, 'Standardowy': 0, 'Premium': 0},
        'packets': {'Miesięczny': 0, 'Rodzinny': 0}
    }
}


def test_set_footer():
    document = Document()
    text = '1'
    utils.set_footer(document, text)

    assert text in document.sections[-1].footer.paragraphs[0].text


def test_create_table_with_valid_size():
    document = Document()
    table = utils.create_table(document, 2, 3)

    assert len(table.rows) == 2
    assert len(table.columns) == 3


def test_create_table_with_invalid_size():
    document = Document()
    with pytest.raises(ValueError):
        utils.create_table(document, -1, -1)


def test_add_nonexistent_image_to_cell():
    document = Document()
    table = utils.create_table(document, 1, 1)
    cell = table.cell(0, 0)
    with pytest.raises(FileNotFoundError):
        utils.add_image_to_cell(cell, 'doesnt_exist.png')


def test_delete_paragraph():
    document = Document()
    document.add_paragraph('example paragraph')
    assert len(document.paragraphs) == 1

    utils.delete_paragraph(document.paragraphs[0])
    assert len(document.paragraphs) == 0


def test_change_cells_text_bold():
    document = Document()
    table = utils.create_table(document, 1, 1)
    utils.change_cells_text_bold([table.cell(0, 0)], True)

    assert table.cell(0, 0).paragraphs[0].runs[1].bold is True


def test_change_cells_text_color():
    document = Document()
    table = utils.create_table(document, 1, 1)
    color = RGBColor(0xFF, 0xFF, 0xFF)
    utils.change_cells_text_color([table.cell(0, 0)], color)

    assert table.cell(0, 0).paragraphs[0].runs[1].font.color.rgb == color


def test_change_cells_background_color():
    document = Document()
    table = utils.create_table(document, 1, 1)
    color = 'FFFFFF'
    utils.change_cells_background_color([table.cell(0, 0)], color)

    assert table.cell(0, 0)._element.tcPr.xpath('w:shd')[0].get(qn('w:fill')) == color


def test_create_stylised_document():
    assert utils.create_stylised_document() is not None


def test_create_sales_table():
    document = Document()
    assert utils.create_sales_table(document, data) is not None


def test_create_packet_summary_table():
    document = Document()
    assert utils.create_packet_summary_table(document, data) is not None


def test_create_subscriptions_summary_table():
    document = Document()
    assert utils.create_subscriptions_summary_table(document, data) is not None


def test_create_summary_table():
    document = Document()
    assert utils.create_summary_table(document, data) is not None


def test_create_company_header():
    document = Document()
    header = utils.create_company_header(document, ['test1'], ['test2'])

    assert header.cell(0, 0).text == 'test1'
    assert header.cell(0, 1).text == 'test2'


def test_create_table_of_contents_with_no_headers():
    document = Document()
    with pytest.raises(ValueError):
        utils.create_table_of_contents(document, [], ['1'])


def test_create_table_of_contents_without_matching_values():
    document = Document()
    with pytest.raises(ValueError):
        utils.create_table_of_contents(document, ['test1', 'test2'], ['1'])


def test_create_table_of_contents_with_no_page_numbers():
    document = Document()
    with pytest.raises(ValueError):
        utils.create_table_of_contents(document, ['test1'], [])


def test_create_table_of_contents():
    document = Document()
    toc = utils.create_table_of_contents(document, ['test1'], ['1'])

    assert toc.cell(0, 0).text == 'test1'
    assert toc.cell(0, 1).text == '1'
