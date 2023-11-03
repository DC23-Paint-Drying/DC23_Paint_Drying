"""
Script generating monthly report
"""


from sys import argv
from os import remove

from docx import Document
from docx.enum import section, text

import report_utils as utils
from report_data import get_report_data


def generate(directory: str = '') -> str:
    """
    Generates .docx report using data from report_data.py get_report_data.

    Args:
        directory:
            Directory to create report file to.
            If empty, creates file in working directory.

    Returns:
        Filepath to generated .docx report or empty string if it fails.

    """

    data = get_report_data()
    filepath = f'{directory}/report-{data["date"]}.docx'

    # generate all chart images
    users_chart = utils.create_donut_chart(list(data["users"]["activity"].values()),
                                           list(data["users"]["activity"].keys()))
    users_age_chart = utils.create_donut_chart(list(data["users"]["age"].values()),
                                               list(data["users"]["age"].keys()))
    users_gender_chart = utils.create_donut_chart(list(data["users"]["gender"].values()),
                                                  list(data["users"]["gender"].keys()))

    try:
        document = Document()
        utils.set_footer(document, 1)

        document.add_heading(f'Raport z {data["date"]}', 0)
        document.add_heading(f'PaintDrying', 1)
        table_of_contents = utils.create_table(document, 1, 2)
        table_of_contents.cell(0, 0).paragraphs[0].text = f'Użytkownicy'
        table_of_contents.cell(0, 0).paragraphs[0].style = f'List Number'
        table_of_contents.cell(0, 0).add_paragraph('Sprzedaż', style='List Number')
        table_of_contents.cell(0, 0).add_paragraph('Ostatni miesiąc', style='List Number')
        table_of_contents.cell(0, 1).paragraphs[0].text = f'2\n3\n3'
        table_of_contents.cell(0, 0).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.LEFT
        table_of_contents.cell(0, 1).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.RIGHT

        document.add_section(section.WD_SECTION.NEW_PAGE)
        utils.set_footer(document, 2)

        document.add_heading(f'Użytkownicy', 2)
        users_table = utils.create_table(document, 1, 3)
        users_table.cell(0, 0).paragraphs[0].text = f'Liczba klientów: {sum(data["users"]["activity"].values())}\n' \
                                                    f'Aktywni klienci: {data["users"]["activity"]["Aktywni"]}\n' \
                                                    f'Nieaktywni klienci: {data["users"]["activity"]["Nieaktywni"]}'
        users_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.LEFT
        users_table.cell(0, 1).merge(users_table.cell(0, 2))
        utils.add_image_to_cell(users_table.cell(0, 1), users_chart)

        users_specs_table = utils.create_table(document, 2, 2)
        utils.add_image_to_cell(users_specs_table.cell(0, 0), users_age_chart)
        utils.add_image_to_cell(users_specs_table.cell(0, 1), users_gender_chart)
        users_specs_table.cell(1, 0).paragraphs[0].text = f'Klienci wg kategorii wiekowych'
        users_specs_table.cell(1, 1).paragraphs[0].text = f'Klienci wg płci'

        document.add_section(section.WD_SECTION.NEW_PAGE)
        utils.set_footer(document, 3)

        document.add_heading(f'Sprzedaż', 2)
        sales_table = utils.create_table(document, 4, 4)
        sales_table.style = 'Table Grid'
        sales_table.cell(1, 0).paragraphs[0].text = f'Liczba abonentów'
        sales_table.cell(2, 0).paragraphs[0].text = f'Przychód'
        sales_table.cell(3, 0).paragraphs[0].text = f'Łączny przychód'
        for index, key in enumerate(data["sales"]):
            sales_table.cell(0, index + 1).paragraphs[0].text = f'{key.capitalize()}'
            sales_table.cell(1, index + 1).paragraphs[0].text = f'{data["sales"][key]["users"]}'
            sales_table.cell(2, index + 1).paragraphs[0].text = f'{data["sales"][key]["profit"]}'
        sales_table.cell(3, 2).merge(sales_table.cell(3, 3))
        sales_table.cell(3, 1).merge(sales_table.cell(3, 2))
        sales_table.cell(3, 1).paragraphs[0].text = f'{sum(map(lambda x: data["sales"][x]["profit"], data["sales"]))}'

        document.add_heading(f'Ostatni miesiąc', 2)
        document.add_paragraph(f'')
        summary = utils.create_table(document, 2, 2)
        summary.cell(0, 0).paragraphs[0].text = f'\U0001F4C8 Nowe abonamenty: ' \
                                                f'{sum(map(lambda x: data["recent"][x]["subscribed"], data["recent"]))}'
        summary.cell(0, 1).paragraphs[0].text = f'\U0001F4C9 Anulowane abonamenty: ' \
                                                f'{sum(map(lambda x: data["recent"][x]["cancelled"], data["recent"]))}'
        summary.cell(0, 0).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.RIGHT
        summary.cell(0, 1).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.LEFT
        summary.cell(1, 0).merge(summary.cell(1, 1))

        summary_table = utils.create_table(summary.cell(1, 0), 3, 4)
        summary_table.style = 'Table Grid'
        summary_table.cell(1, 0).paragraphs[0].text = f'Zakupione abonamenty'
        summary_table.cell(2, 0).paragraphs[0].text = f'Anulowane abonamenty'
        for index, key in enumerate(data["recent"]):
            summary_table.cell(0, index + 1).paragraphs[0].text = f'{key.capitalize()}'
            summary_table.cell(1, index + 1).paragraphs[0].text = f'{data["recent"][key]["subscribed"]}'
            summary_table.cell(2, index + 1).paragraphs[0].text = f'{data["recent"][key]["cancelled"]}'

        document.save(f'{directory}report-{data["date"]}.docx')

    except Exception:
        filepath = f''

    finally:
        # image cleanup
        remove(users_chart)
        remove(users_age_chart)
        remove(users_gender_chart)

    return filepath


# for development and command line generation
if __name__ == '__main__':
    globals()[argv[1]](argv[2] if 2 < len(argv) else './')
