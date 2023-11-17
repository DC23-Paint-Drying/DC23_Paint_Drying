"""
Module containing report generation.
"""

from os import remove
from os.path import join, normpath
from sys import argv

from docx.enum import section

import src.report.report_charts as charts
from src.report.report_data import get_report_data
import src.report.report_utils as utils
from src.database_context import DatabaseContext
import src.manifest as company


def generate(db: DatabaseContext, directory: str = '') -> str:
    """
    Generates .docx report using data from report_data.py get_report_data.
    Will throw error if anything fails during report generation.

    Args:
        db:
            Database to get data from.
        directory:
            Directory to create report file to.
            If empty, creates file in working directory.

    Returns:
        Filepath to generated .docx report.
    """
    data = get_report_data(db)
    directory = normpath(join(directory, ''))
    filepath = normpath(join(directory, f'report-{data["date"]}.docx'))

    # generate all chart images
    users_chart = charts.create_donut_chart(list(data["users"]["subscriptions"].values()),
                                            list(data["users"]["subscriptions"].keys()))
    users_age_chart = charts.create_donut_chart(list(data["users"]["age"].values()),
                                                list(data["users"]["age"].keys()))
    users_gender_chart = charts.create_donut_chart(list(data["users"]["gender"].values()),
                                                   list(data["users"]["gender"].keys()))

    try:
        # creates blank document with styles
        document = utils.create_stylised_document()
        utils.set_footer(document, '1')

        # set first page header
        company_header_left = [f'{company.COMPANY_NAME}', f'{company.COMPANY_ADDRESS}', f'NIP {company.COMPANY_NIP}']
        company_header_right = [f'{data["timestamp"]}']
        utils.create_company_header(document, company_header_left, company_header_right)

        # set title and toc
        document.add_heading(f'Raport miesięczny', 1)
        document.add_heading(f'Okres {data["date"]}', 2)
        utils.create_table_of_contents(document, ['Użytkownicy', 'Sprzedaż', 'Ostatni miesiąc'], ['2', '3', '3'])

        document.add_section(section.WD_SECTION.NEW_PAGE)
        utils.set_footer(document, '2')

        # add table containing users data
        document.add_heading(f'Użytkownicy', 3)
        users_table = utils.create_users_statistics_table(document, data, company)
        utils.add_image_to_cell(users_table.cell(0, 1), users_chart)
        document.add_heading(f'', 2)

        # add table containing graphs about users
        users_specs_table = utils.create_table(document, 2, 2)
        utils.add_image_to_cell(users_specs_table.cell(0, 0), users_age_chart)
        utils.add_image_to_cell(users_specs_table.cell(0, 1), users_gender_chart)
        users_specs_table.cell(1, 0).paragraphs[0].text = f'Klienci wg kategorii wiekowych'
        users_specs_table.cell(1, 1).paragraphs[0].text = f'Klienci wg płci'

        document.add_section(section.WD_SECTION.NEW_PAGE)
        utils.set_footer(document, '3')

        # add sales section and sales table
        document.add_heading(f'Sprzedaż', 3)
        utils.create_sales_table(document, data)
        document.add_heading(f'', 4)

        # add recent section
        document.add_heading(f'Ostatni miesiąc', 3)
        summary = utils.create_summary_table(document, data)

        # add recent subscriptions table
        utils.delete_paragraph(summary.cell(1, 0).paragraphs[0])
        utils.create_subscriptions_summary_table(summary.cell(1, 0), data)

        # add recent packets table
        utils.create_packet_summary_table(summary.cell(1, 0), data)

        document.save(filepath)

    except Exception:
        filepath = f''

    finally:
        # image cleanup
        remove(users_chart)
        remove(users_age_chart)
        remove(users_gender_chart)

    if not filepath:
        raise Exception('Failed to generate report.')

    return filepath


# for development and command line generation
if __name__ == '__main__':
    db = DatabaseContext('db')
    globals()[argv[1]](db, argv[2] if 2 < len(argv) else './')
