"""
Module containing utility functions used during report generation.
"""


from subprocess import check_output
from os import environ, getcwd
from os.path import splitext, basename, join

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.enum import table, text
from docx.shared import Emu, Pt, RGBColor
from docx.oxml.shared import qn, OxmlElement


LIBREOFFICE = environ.get("LIBREOFFICE", "")
if not LIBREOFFICE:
    print("Environmental variable LIBREOFFICE not set. PDF convertion won't function!")
    print("LIBREOFFICE should be set to path to soffice/soffice.exe included with LibreOffice installation")


def create_stylised_document() -> Document:
    """
    Creates document and sets all styles for headings and tables.

    Returns:
        Created document.
    """
    document = Document()

    # color definitions
    black = RGBColor(0x23, 0x23, 0x23)
    gray = RGBColor(0x6a, 0x6a, 0x6a)

    # non heading text
    normal = document.styles['Normal']
    normal.font.size = Pt(13)
    normal.font.bold = False
    normal.font.italic = True
    normal.font.name = 'Calibri'
    normal.font.color.rgb = black

    # title
    heading1 = document.styles['Heading 1']
    heading1.paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.space_before = Pt(26)
    heading1.font.size = Pt(26)
    heading1.font.bold = False
    heading1.font.italic = False
    heading1.font.name = 'Arial'
    heading1.font.color.rgb = black

    # subtitle
    heading2 = document.styles['Heading 2']
    heading2.paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.CENTER
    heading2.paragraph_format.space_before = False
    heading2.paragraph_format.space_after = Pt(10)
    heading2.font.size = Pt(15)
    heading2.font.bold = False
    heading2.font.italic = False
    heading2.font.name = 'Arial'
    heading2.font.color.rgb = gray

    # section heading
    heading3 = document.styles['Heading 3']
    heading3.paragraph_format.space_before = Pt(16)
    heading3.paragraph_format.space_after = Pt(6)
    heading3.font.size = Pt(20)
    heading3.font.bold = True
    heading3.font.italic = False
    heading3.font.name = 'Calibri'
    heading3.font.color.rgb = black

    # table of contents
    heading4 = document.styles['Heading 4']
    heading4.paragraph_format.space_before = False
    heading4.paragraph_format.space_after = False
    heading4.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.SINGLE
    heading4.font.size = Pt(14)
    heading4.font.bold = True
    heading4.font.italic = False
    heading4.font.name = 'Calibri'
    heading4.font.color.rgb = black

    # page number
    heading5 = document.styles['Heading 5']
    heading5.paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.RIGHT
    heading5.paragraph_format.space_before = False
    heading5.paragraph_format.space_after = False
    heading5.font.size = Pt(12)
    heading5.font.bold = False
    heading5.font.italic = False
    heading5.font.name = 'Calibri'
    heading5.font.color.rgb = black

    # company header
    heading6 = document.styles['Heading 6']
    heading6.paragraph_format.space_before = False
    heading6.paragraph_format.space_after = False
    heading6.paragraph_format.line_spacing_rule = text.WD_LINE_SPACING.SINGLE
    heading6.font.size = Pt(14)
    heading6.font.bold = False
    heading6.font.italic = False
    heading6.font.name = 'Calibri'
    heading6.font.color.rgb = black

    return document


def create_users_statistics_table(document, data, company_data):
    """
    Adds table containing data regarding users.

    Args:
        document:
            Document to add table to.
        data:
            Data about users.
        company_data:
            Data about company.

    Returns:
        Created table.
    """
    users_table = create_table(document, 1, 3)
    users_table.autofit = False
    users_table.cell(0, 0).paragraphs[0].add_run(f'Liczba klientów: ')
    users_table.cell(0, 0).paragraphs[0].add_run(f'{sum(data["users"]["subscriptions"].values())}\n').bold = True
    users_table.cell(0, 0).paragraphs[0].add_run(f'Dostępne abonamenty: ')
    users_table.cell(0, 0).paragraphs[0].add_run(f'{len(company_data.SUBSCRIPTIONS)}\n').bold = True
    users_table.cell(0, 0).paragraphs[0].add_run(f'Dostępne pakiety: ')
    users_table.cell(0, 0).paragraphs[0].add_run(f'{len(company_data.PACKETS)}\n').bold = True
    users_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.LEFT
    users_table.cell(0, 1).merge(users_table.cell(0, 2))

    return users_table


def create_sales_table(document, data):
    """
    Adds table containing data regarding sales.

    Args:
        document:
            Document to add table to.
        data:
            Data about sales.

    Returns:
        Created table.
    """
    sales_table = create_table(document, 4, 4)
    sales_table.style = 'Table Grid'
    sales_table.cell(1, 0).paragraphs[0].text = f'Liczba abonentów'
    sales_table.cell(2, 0).paragraphs[0].text = f'Przychód'
    sales_table.cell(3, 0).paragraphs[0].text = f'Łączny przychód'
    for index, key in enumerate(data["sales"]):
        sales_table.cell(0, index + 1).paragraphs[0].text = f'{key.capitalize()}'
        sales_table.cell(1, index + 1).paragraphs[0].text = f'{data["sales"][key]["users"]}'
        sales_table.cell(2, index + 1).paragraphs[0].text = f'{data["sales"][key]["profit"]:.2f} zł'
    sales_table.cell(3, 2).merge(sales_table.cell(3, 3))
    sales_table.cell(3, 1).merge(sales_table.cell(3, 2))
    sales_table.cell(3, 1).paragraphs[
        0].text = f'{sum(map(lambda x: data["sales"][x]["profit"], data["sales"])):.2f} zł'
    change_cells_text_bold(sales_table.rows[0].cells[1:], True)
    change_cells_text_bold(sales_table.rows[3].cells, True)
    change_cells_background_color(sales_table.rows[0].cells, 'B7B7B7')
    change_cells_background_color(sales_table.rows[1].cells, 'EFEFEF')
    change_cells_background_color(sales_table.rows[2].cells, 'B7B7B7')
    change_cells_background_color(sales_table.rows[3].cells, 'EFEFEF')

    return sales_table


def create_packet_summary_table(document, data):
    """
    Adds table containing data regarding packets.

    Args:
        document:
            Document to add table to.
        data:
            Data about packets.

    Returns:
        Created table.
    """
    summary_packets = create_table(document, 2, 1 + len(data["recent"]["packets"]))
    summary_packets.style = 'Table Grid'
    summary_packets.cell(1, 0).paragraphs[0].text = f'Zakupione pakiety'
    for index, key in enumerate(data["recent"]["packets"]):
        summary_packets.cell(0, index + 1).paragraphs[0].text = f'{key.capitalize()}'
        summary_packets.cell(1, index + 1).paragraphs[0].text = f'{data["recent"]["packets"][key]}'
    change_cells_text_bold(summary_packets.rows[0].cells[1:], True)
    change_cells_background_color(summary_packets.rows[0].cells, 'B7B7B7')
    change_cells_background_color(summary_packets.rows[1].cells, 'EFEFEF')

    return summary_packets


def create_subscriptions_summary_table(document, data):
    """
    Adds table containing data regarding subscriptions sales.

    Args:
        document:
            Document to add table to.
        data:
            Data about subscriptions sales.

    Returns:
        Created table.
    """
    summary_subscriptions = create_table(document, 2, 1 + len(data["recent"]["subscribed"]))
    summary_subscriptions.style = 'Table Grid'
    summary_subscriptions.cell(1, 0).paragraphs[0].text = f'Zakupione abonamenty'
    for index, key in enumerate(data["recent"]["subscribed"]):
        summary_subscriptions.cell(0, index + 1).paragraphs[0].text = f'{key.capitalize()}'
        summary_subscriptions.cell(1, index + 1).paragraphs[0].text = f'{data["recent"]["subscribed"][key]}'
    change_cells_text_bold(summary_subscriptions.rows[0].cells[1:], True)
    change_cells_background_color(summary_subscriptions.rows[0].cells, 'B7B7B7')
    change_cells_background_color(summary_subscriptions.rows[1].cells, 'EFEFEF')

    return summary_subscriptions


def create_summary_table(document, data):
    """
    Adds table containing data regarding recent trends.

    Args:
        document:
            Document to add table to.
        data:
            Data about recent trends.

    Returns:
        Created table.
    """
    summary = create_table(document, 2, 3)
    summary.cell(1, 1).merge(summary.cell(1, 2))
    summary.cell(1, 0).merge(summary.cell(1, 1))

    summary_left = summary.cell(0, 0).paragraphs[0]
    summary_left.text = f'Nowi użytkownicy: '
    sl = summary_left.add_run(f'{data["recent"]["users"]}')
    sl.bold = True
    sl.font.color.rgb = RGBColor(0x38, 0x76, 0x1D)
    summary_left.paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.RIGHT

    summary_center = summary.cell(0, 1).paragraphs[0]
    summary_center.text = f'Nowe abonamenty: '
    sc = summary_center.add_run(f'{sum(map(lambda x: data["recent"]["subscribed"][x], data["recent"]["subscribed"]))}')
    sc.bold = True
    sc.font.color.rgb = RGBColor(0x38, 0x76, 0x1D)
    summary_center.paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.CENTER

    summary_right = summary.cell(0, 2).paragraphs[0]
    summary_right.text = f'Nowe pakiety: '
    sr = summary_right.add_run(f'{sum(map(lambda x: data["recent"]["packets"][x], data["recent"]["packets"]))}')
    sr.bold = True
    sr.font.color.rgb = RGBColor(0x38, 0x76, 0x1D)
    summary_right.paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.LEFT

    return summary


def create_company_header(document: Document, left: [str], right: [str]) -> Table:
    """
    Adds table containing company information in 2 column layout.

    Args:
        document:
            Document to add table to.
        left:
            List of information to be displayed on the left of the page.
        right:
            List of information to be displayed on the right of the page.

    Returns:
        Created table.
    """

    company_table = create_table(document, 1, 2)
    company_table.cell(0, 0).paragraphs[0].text = f'\n'.join(left)
    company_table.cell(0, 0).paragraphs[0].style = f'Heading 6'
    company_table.cell(0, 0).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.LEFT
    company_table.cell(0, 0).vertical_alignment = table.WD_CELL_VERTICAL_ALIGNMENT.TOP
    company_table.cell(0, 1).paragraphs[0].text = f'\n'.join(right)
    company_table.cell(0, 1).paragraphs[0].style = f'Heading 6'
    company_table.cell(0, 1).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.RIGHT
    company_table.cell(0, 1).vertical_alignment = table.WD_CELL_VERTICAL_ALIGNMENT.TOP

    return company_table


def create_table_of_contents(document: Document, headers: [str], page_numbers: [str]) -> Table:
    """
    Creates table of contents table with provided headers and their page numbers.

    Args:
        document:
            Document to add toc to.
        headers:
            List of headers for toc.
        page_numbers:
            List of page numbers for each header.

    Returns:
        Created table.
    """

    if not headers or not page_numbers:
        raise ValueError('Values and labels must not be empty')

    if len(headers) != len(page_numbers):
        raise ValueError('Each value must have exactly one label')

    toc = create_table(document, 1, 2)
    toc.cell(0, 0).paragraphs[0].text = f'\n'.join(headers)
    toc.cell(0, 0).paragraphs[0].style = f'Heading 4'
    toc.cell(0, 0).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.LEFT
    toc.cell(0, 1).paragraphs[0].text = f'\n'.join(page_numbers)
    toc.cell(0, 1).paragraphs[0].style = f'Heading 4'
    toc.cell(0, 1).paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.RIGHT

    return toc


def set_footer(document: Document, page_number: str) -> None:
    """
    Sets page number for current section.

    Args:
        document:
            Document to add footer to.
        page_number:
            Page number to add in the right corner of footer.

    Returns:
        None
    """
    document.sections[-1].footer.is_linked_to_previous = False
    document.sections[-1].footer.paragraphs[0].text = f'\t\t{page_number}'
    document.sections[-1].footer.paragraphs[0].style = 'Heading 5'


def create_table(document: Document, rows: int, cols: int) -> Table:
    """
    Creates and adds new table of size rows x cols.
    The table is centered on page.
    All cells are vertically and horizontally aligned.

    Args:
        document:
            Document to add table to.
        rows:
            Amount of rows in table.
        cols:
            Amount of columns in table.

    Returns:
        Created table.
    """

    if rows < 1 or cols < 1:
        raise ValueError('Row and column size must be atleast 1')

    t = document.add_table(rows=rows, cols=cols)
    t.alignment = table.WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        row.height_rule = table.WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Pt(30)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = text.WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = table.WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return t


def delete_paragraph(paragraph: Paragraph) -> None:
    """
    Deletes passed paragraph from parent element.

    Args:
        paragraph:
            Paragraph to delete.

    Returns:
        None
    """
    p = paragraph._element
    p.getparent().remove(p)
    p._p = None
    p._element = None


def add_image_to_cell(cell: _Cell, image: str) -> None:
    """
    Adds image to provided table cell.

    Args:
        cell:
            Reference to cell.
        image:
            File path to image.

    Returns:
        None
    """
    # image width must be smaller than cell width to accomodate pdf conversion inacurracies
    cell.paragraphs[0].add_run().add_picture(image, Emu(cell.width * 0.85))


def change_cells_text_bold(cells: [_Cell], value: bool) -> None:
    """
    Changes text bold setting for each cell in list.

    Args:
        cells:
            List of cells to change text of.
        value:
            Bold value to set text to.

    Returns:
        None
    """
    for cell in cells:
        t = cell.paragraphs[0].text
        cell.paragraphs[0].text = f''
        cell.paragraphs[0].add_run(t).bold = value


def change_cells_text_color(cells: [_Cell], color: RGBColor) -> None:
    """
    Changes text color setting for each cell in list.

    Args:
        cells:
            List of cells to change text of.
        color:
            RGB color to set text to.

    Returns:
        None
    """
    for cell in cells:
        t = cell.paragraphs[0].text
        cell.paragraphs[0].text = f''
        cell.paragraphs[0].add_run(t).font.color.rgb = color


def change_cells_background_color(cells: [_Cell], color: str) -> None:
    """
    Changes background color setting for each cell in list.

    Args:
        cells:
            List of cells to change background of.
        color:
            RGB color to set background to.

    Returns:
        None
    """
    for cell in cells:
        properties = cell._element.tcPr
        try:
            shading = properties.xpath('w:shd')[0]
        except IndexError:
            shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        properties.append(shading)


def convert_docx_to_pdf(docx_path: str, pdf_directory: str) -> str:
    """
    Calls LibreOffice to convert provided docx_file to pdf.
    Converted pdf file is saved in pdf_directory.
    LibreOffice soffice utility is provided via environmental variable LIBREOFFICE.

    Args:
        docx_path:
            Path to docx file to convert.
        pdf_directory:
            Directory to save pdf file in.

    Returns:
        Path to pdf file.
    """

    if not LIBREOFFICE:
        raise RuntimeError('LIBREOFFICE environmental variable is not set, which is required to convert to pdf.')

    args = [str(LIBREOFFICE), '--headless', '--convert-to', 'pdf', str(docx_path), '--outdir', str(pdf_directory)]
    check_output(args, cwd=getcwd(), timeout=20)

    return join(pdf_directory, splitext(basename(docx_path))[0] + '.pdf')
